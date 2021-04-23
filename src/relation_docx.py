#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Time     : 2021-04-15 16:49
# @Author   : San
# @Project  : Malaysian-Character-Knowledge-Graph
# @File     : relation_docx.py
# @Software : PyCharm
# @Desc     :


from docx import Document
from docx.shared import Cm
import json

if __name__ == '__main__':
    with open('../ms_wiki_data/relations.json', 'r', encoding='utf-8') as f:
        rel = json.load(f)
    style = 'Table Grid'

    document = Document()

    table = document.add_table(len(rel) + 1, 3)
    table.style = style
    table.rows[0].cells[0].text = '关系名（en）'
    table.rows[0].cells[0].width = Cm(5)
    table.rows[0].cells[1].text = '关系名（zh）'
    table.rows[0].cells[1].width = Cm(5)
    table.rows[0].cells[2].text = '内部关系'
    table.rows[0].cells[2].width = Cm(15)

    for i, j in enumerate(rel):
        table.rows[i + 1].cells[0].text = j
        table.rows[i + 1].cells[0].width = Cm(5)
        table.rows[i + 1].cells[1].text = rel[j]['zh']
        table.rows[i + 1].cells[1].width = Cm(5)
        table.rows[i + 1].cells[2].width = Cm(15)
        if j.__contains__('multi'):
            i_rel = rel[j]['inner_relation']
            i_table = table.rows[i + 1].cells[2].add_table(len(i_rel) + 1, 2)
            i_table.style = style
            i_table.rows[0].cells[0].text = '关系名（en）'
            i_table.rows[0].cells[1].text = '关系名（zh）'
            for ii, jj in enumerate(i_rel):
                i_table.rows[ii + 1].cells[0].text = jj
                i_table.rows[ii + 1].cells[1].text = i_rel[jj]['zh']

    document.save('../ms_wiki_data/relation.docx')
