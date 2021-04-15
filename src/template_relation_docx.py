#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Time     : 2021-04-15 18:10
# @Author   : San
# @Project  : Malaysian-Character-Knowledge-Graph
# @File     : template_relation_docx.py
# @Software : PyCharm
# @Desc     :


from docx import Document
from docx.shared import Cm
import json

if __name__ == '__main__':
    with open('../ms_wiki_data/template_relations.json', 'r', encoding='utf-8') as f:
        rel = json.load(f)
    style = 'Table Grid'

    document = Document()
    table = document.add_table(len(rel) + 1, 3)
    table.style = style
    table.rows[0].cells[0].text = '模板名'
    table.rows[0].cells[0].width = Cm(5)
    table.rows[0].cells[1].text = '最显著职业'
    table.rows[0].cells[1].width = Cm(5)
    table.rows[0].cells[2].text = '关系'
    table.rows[0].cells[2].width = Cm(15)

    for i, j in enumerate(rel):
        table.rows[i + 1].cells[0].text = 'Template' + j.replace(' ', '')
        table.rows[i + 1].cells[0].width = Cm(5)
        table.rows[i + 1].cells[1].text = j
        table.rows[i + 1].cells[1].width = Cm(5)
        table.rows[i + 1].cells[2].width = Cm(15)
        l_rel = rel[j]
        l_table = table.rows[i + 1].cells[2].add_table(len(l_rel) + 1, 3)
        l_table.style = style
        l_table.rows[0].cells[0].text = '关系名（en）'
        l_table.rows[0].cells[1].text = '关系名（zh）'
        l_table.rows[0].cells[2].text = '内部关系'
        for ii, jj in enumerate(l_rel):
            l_table.rows[ii + 1].cells[0].text = jj
            l_table.rows[ii + 1].cells[1].text = l_rel[jj]['zh']
            if jj.__contains__('multi'):
                i_rel = l_rel[jj]['inner_relation']
                i_table = l_table.rows[ii + 1].cells[2].add_table(len(i_rel) + 1, 2)
                i_table.style = style
                i_table.rows[0].cells[0].text = '关系名（en）'
                i_table.rows[0].cells[1].text = '关系名（zh）'
                for iii, jjj in enumerate(i_rel):
                    i_table.rows[iii + 1].cells[0].text = jjj
                    i_table.rows[iii + 1].cells[1].text = i_rel[jjj]['zh']

    document.save('template_relation.docx')
